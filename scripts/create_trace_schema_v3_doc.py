from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "documents" / "Trace Schema Definitons-v2.docx"
OUTPUT_DOC = ROOT / "documents" / "Trace Schema Definitons-v3.docx"


def set_cell(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(8)


def add_table_row(table, values):
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)
    return row


def fill_row(row, values):
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)


def add_small_table_before(doc, target_paragraph, rows):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    fill_row(table.rows[0], rows[0])
    for value_row in rows[1:]:
        add_table_row(table, value_row)
    target_paragraph._p.addprevious(table._tbl)
    return table


def add_paragraph_before(target_paragraph, text="", style=None):
    paragraph = target_paragraph.insert_paragraph_before(text)
    if style:
        paragraph.style = style
    return paragraph


def add_code_block_before(target_paragraph, text):
    paragraph = add_paragraph_before(target_paragraph)
    paragraph.style = "No Spacing"
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    return paragraph


def find_paragraph(doc, text):
    for paragraph in doc.paragraphs:
        if text in paragraph.text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def replace_paragraph_text(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def clear_table_after_header(table):
    for row in table.rows[1:]:
        for cell in row.cells:
            set_cell(cell, "")


def extend_table_to_rows(table, row_count):
    while len(table.rows) < row_count:
        table.add_row()


def update_existing_content(doc):
    replace_paragraph_text(find_paragraph(doc, "XML Schema Definiton"), "XML Schema Definiton v3")
    replace_paragraph_text(
        find_paragraph(doc, "Bu dokümantasyon iki şema versiyonunu kapsamaktadır:"),
        "Bu dokümantasyon v1.0'dan v2.3'e kadar olan şema sürümlerini kapsamaktadır:",
    )
    replace_paragraph_text(find_paragraph(doc, '<trace schema_version="2.0">'), '<trace schema_version="2.3">')
    insert_after = find_paragraph(doc, "<events>")
    insert_after.insert_paragraph_before("<!-- Program ve çalışma ortamı bilgileri -->")

    meta = doc.tables[0]
    set_cell(meta.rows[0].cells[1], "26.07.2026")
    set_cell(meta.rows[3].cells[1], "v3 (Trace Schema v2.3)")

    versions = doc.tables[2]
    extend_table_to_rows(versions, 6)
    fill_row(versions.rows[0], ["Versiyon", "Açıklama", "Not"])
    rows = [
        ["v2.3", "Dynamic slicing desteği: event içinde <reads>, trace altında opsiyonel <slice>, SliceNodeType ve SliceDependencyEnum.", "Yeni sürüm"],
        ["v2.2", "Scope/iteration zamanlama attribute'ları, exception traceback, metadata targets/source_files.", "Geriye uyumlu"],
        ["v2.1", "IterationType içinde scope/loop nesting desteği ve DeltaChangeType variable attribute'u.", "Ara sürüm"],
        ["v2.0", "Hiyerarşik yapı: <scope>, <loop>, <thread>, delta-only line event yaklaşımı.", "Ana kırılım"],
        ["v1.0", "Flat event listesi; her event içinde tam locals verisi.", "İlk sürüm"],
    ]
    for row, values in zip(versions.rows[1:], rows):
        fill_row(row, values)

    add_table_row(doc.tables[3], ["<slice>", "Backward dynamic slice sonucunu taşıyan opsiyonel kök seviye element. Bkz. §2.16", "Hayır"])
    add_table_row(doc.tables[4], ["<targets>", "Trace sırasında hedef alınan fonksiyon ve dosyalar (v2.2).", "Hayır"])
    add_table_row(doc.tables[4], ["<source_files>", "Event kaydedilen kaynak dosyaların listesi (v2.2).", "Hayır"])
    add_table_row(doc.tables[8], ["<reads>", "Satırın okuduğu/kullandığı değişkenler. Dynamic slicing use-kümesini taşır (v2.3).", "Hayır"])
    add_table_row(doc.tables[9], ["start_time (attribute)", "Fonksiyon çağrısının UNIX timestamp'i (v2.2).", "Hayır"])
    add_table_row(doc.tables[9], ["end_time (attribute)", "Fonksiyonun return/exception UNIX timestamp'i (v2.2).", "Hayır"])
    add_table_row(doc.tables[9], ["duration (attribute)", "Fonksiyon süresi, saniye cinsinden (end_time - start_time) (v2.2).", "Hayır"])
    add_table_row(doc.tables[13], ["<scope>", "İterasyon içinde çağrılan fonksiyon scope'u (v2.1).", "Hayır"])
    add_table_row(doc.tables[13], ["<loop>", "İterasyon içindeki nested loop (v2.1).", "Hayır"])
    add_table_row(doc.tables[13], ["start_time (attribute)", "İterasyon başlangıç UNIX timestamp'i (v2.2).", "Hayır"])
    add_table_row(doc.tables[13], ["end_time (attribute)", "İterasyon bitiş UNIX timestamp'i (v2.2).", "Hayır"])
    add_table_row(doc.tables[20], ["<traceback>", "Exception'ın tam stack trace metni. Post-mortem debugging için kullanılır (v2.2).", "Hayır"])

    quick = doc.tables[22]
    quick.add_column(Inches(0.7))
    quick.add_column(Inches(0.7))
    fill_row(quick.rows[0], ["Element", "Kısa Açıklama", "v1.0", "v2.0-v2.2", "v2.3"])
    for row in quick.rows[1:]:
        existing = [cell.text for cell in row.cells]
        fill_row(row, existing[:3] + [existing[3] if len(existing) > 3 else "", existing[3] if len(existing) > 3 else ""])
    quick_rows = [
        ["<targets>", "Trace hedef fonksiyon/dosya listesi", "✗", "✓", "✓"],
        ["<source_files>", "Trace edilen kaynak dosya listesi", "✗", "✓", "✓"],
        ["<traceback>", "Exception stack trace metni", "✗", "✓", "✓"],
        ["<reads>", "Line event'in okuduğu değişkenler", "✗", "✗", "✓"],
        ["<slice>", "Backward dynamic slice sonucu", "✗", "✗", "✓"],
        ["<node>", "Slice'a giren event/sentetik düğüm", "✗", "✗", "✓"],
        ["SliceDependencyEnum", "criterion/data/control bağımlılık tipi", "✗", "✗", "✓"],
    ]
    for values in quick_rows:
        add_table_row(quick, values)


def insert_v23_reference_sections(doc):
    target = find_paragraph(doc, "3. Örnek XML Dosyaları")

    add_paragraph_before(target, "2.15  <reads>  - Okunan Değişkenler", "Heading 2")
    add_paragraph_before(
        target,
        "v2.3 ile EventType altına eklenen <reads> elementi, özellikle line event'lerinde satırın okuduğu değişkenleri kaydeder. Bu bilgi backward dynamic slicing için use-kümesini oluşturur. Her <read> öğesi ValueType yapısını kullanır; name ve type attribute'ları ile değer metnini taşır.",
    )
    add_small_table_before(
        doc,
        target,
        [
            ["Özellik / Element", "Açıklama", "Zorunlu?"],
            ["<read>", "Okunan tek değişken. ValueType olduğu için name, type ve text değerini içerir.", "Hayır"],
        ],
    )
    add_code_block_before(
        target,
        '<reads>\n  <read name="a" type="int">3</read>\n  <read name="b" type="int">5</read>\n</reads>',
    )

    add_paragraph_before(target, "2.16  <slice>  - Backward Dynamic Slice", "Heading 2")
    add_paragraph_before(
        target,
        "v2.3 ile trace kök elementinin sonuna opsiyonel <slice> bloğu eklenebilir. Bu blok, seçilen kriterden geriye doğru hesaplanan dynamic slice sonucunu taşır. Böylece LLM veya analiz aracı, tüm trace yerine sadece kriteri etkileyen event/düğüm alt kümesini tüketebilir.",
    )
    add_small_table_before(
        doc,
        target,
        [
            ["Özellik / Attribute", "Açıklama", "Zorunlu?"],
            ["criterion", "Slice kriteri. Değişken adı, ifade veya kullanıcı tanımlı kriter metni olabilir.", "Evet"],
            ["criterion_event", "Kriterin dayandığı event id'si.", "Hayır"],
            ["include_control", "Control dependency düğümlerinin dahil edilip edilmediği.", "Hayır"],
            ["<node>", "Slice'a dahil edilen event veya sentetik düğüm.", "Hayır"],
        ],
    )

    add_paragraph_before(target, "2.17  <node>  - Slice Düğümü", "Heading 2")
    add_paragraph_before(
        target,
        "Her <node>, slice sonucuna giren bir satırı temsil eder. line zorunludur; event_id opsiyoneldir çünkü sentetik loop-header düğümlerinin birebir event karşılığı olmayabilir.",
    )
    add_small_table_before(
        doc,
        target,
        [
            ["Attribute", "Tip", "Zorunlu?"],
            ["line", "xs:nonNegativeInteger", "Evet"],
            ["event_id", "xs:nonNegativeInteger", "Hayır"],
            ["function", "xs:string", "Hayır"],
            ["file", "xs:string", "Hayır"],
            ["dependency", "SliceDependencyEnum", "Hayır"],
            ["source", "xs:string", "Hayır"],
        ],
    )

    add_paragraph_before(target, "2.18  SliceDependencyEnum", "Heading 2")
    add_paragraph_before(
        target,
        "Slice düğümünün neden sonuca dahil edildiğini belirtir: criterion kriter event'ini, data veri bağımlılığını, control ise kontrol bağımlılığını ifade eder.",
    )
    add_small_table_before(
        doc,
        target,
        [
            ["Değer", "Açıklama"],
            ["criterion", "Slice kriterinin kendisi"],
            ["data", "Veri bağımlılığı nedeniyle dahil edilen düğüm"],
            ["control", "Kontrol bağımlılığı nedeniyle dahil edilen düğüm"],
        ],
    )


def insert_v23_example(doc):
    target = find_paragraph(doc, "4. v1.0 Şema Tanımı")
    add_paragraph_before(target, "3.6 Dynamic Slicing Örneği - v2.3", "Heading 2")
    add_paragraph_before(
        target,
        "Aşağıdaki örnek, result değişkeni için hesaplanan slice çıktısında hem satırın okuduğu değişkenlerin hem de slice sonucunun nasıl temsil edildiğini gösterir.",
    )
    add_code_block_before(
        target,
        '''<trace schema_version="2.3">
  <metadata>...</metadata>
  <events>
    <scope function="add" file="example.py" call_line="1" depth="1">
      <event id="2" type="line" timestamp="1705312800.002" depth="1">
        <line>2</line>
        <source>result = a + b</source>
        <delta>
          <change name="result" action="added" type="int"><new>8</new></change>
        </delta>
        <reads>
          <read name="a" type="int">3</read>
          <read name="b" type="int">5</read>
        </reads>
      </event>
    </scope>
  </events>
  <slice criterion="result" criterion_event="2" include_control="true">
    <node line="2" event_id="2" function="add" file="example.py"
          dependency="criterion" source="result = a + b"/>
  </slice>
</trace>''',
    )


def normalize_title_and_header(doc):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in {"Python Execution Trace", "XML Schema Definiton v3"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True


def main():
    doc = Document(SOURCE_DOC)
    update_existing_content(doc)
    insert_v23_reference_sections(doc)
    insert_v23_example(doc)
    normalize_title_and_header(doc)
    doc.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
